import streamlit as st
import base64
import tempfile
import os
import io
import re
import google.generativeai as genai
from docx import Document
import datetime
import json
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import TypedDict
from langgraph.graph import StateGraph, END
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import HumanMessage

# ==========================================
# 1. Define the LangGraph State
# ==========================================
class PaperState(TypedDict):
    api_key: str
    marks: int
    subject: str
    grade_class: str
    school_name: str
    session: str
    exam_time: str
    board_format: bool
    is_ncert: bool
    ncert_chapters: str
    comments: str
    file_data: list  # Store base64 encoded files with mime types
    generated_paper: str
    generated_answer_key: str

# ==========================================
# 2. Define Graph Nodes
# ==========================================
def generate_question_paper(state: PaperState):
    """Node that calls Gemini to generate the paper."""
    
    # Initialize the LLM with the user's key
    llm = ChatGoogleGenerativeAI(
        model="gemini-3.6-flash", 
        api_key=state["api_key"]
    )
    
    # Construct the prompt text
    prompt_text = f"""
    You are an expert teacher. Generate a question paper strictly based on the provided images (if any) and the following criteria:
    - Subject: {state['subject']}
    - Class/Grade: {state['grade_class']}
    - Total Marks: {state['marks']}
    - Question Type: Mixed / Balanced (include MCQs, short answers, and long answers)
    """
    
    if state.get("board_format"):
        prompt_text += "\n    - Structure: Strictly follow the official Haryana Board (HBSE) Exam Format (e.g., proper sections, choice of questions, instructions at the top)."
    
    if state.get("is_ncert"):
        prompt_text += f"\n    - Syllabus/Curriculum: Strictly follow NCERT guidelines."
        if state.get("ncert_chapters"):
            prompt_text += f" Focus on the following chapters: {state['ncert_chapters']}"
            
    prompt_text += f"""
    - Additional Instructions: {state['comments']}
    
    Format the output clearly with sections, question numbers, and marks per question.
    """
    
    # Build the message content handling both text and multiple images
    message_content = [{"type": "text", "text": prompt_text}]
    
    for file_info in state.get("file_data", []):
        mime_type = file_info["mime_type"]
        
        if mime_type == "application/pdf":
            if "file_uri" in file_info:
                message_content.append({
                    "type": "media",
                    "mime_type": mime_type,
                    "file_uri": file_info["file_uri"]
                })
            else:
                message_content.append({
                    "type": "media",
                    "mime_type": mime_type,
                    "data": file_info.get("data", "")
                })
        elif mime_type.startswith("image/"):
            b64_data = file_info["data"]
            message_content.append({
                "type": "image_url",
                "image_url": {"url": f"data:{mime_type};base64,{b64_data}"}
            })
        
    message = HumanMessage(content=message_content)
    
    # Generate the response
    response = llm.invoke([message])
    
    content = response.content
    if isinstance(content, list):
        text_parts = []
        for item in content:
            if isinstance(item, dict) and "text" in item:
                text_parts.append(item["text"])
            elif isinstance(item, str):
                text_parts.append(item)
        text_content = "\n".join(text_parts)
    else:
        text_content = str(content)
        
    return {"generated_paper": text_content}

def generate_answer_key(state: PaperState):
    """Node that calls Gemini to generate the answer key for the generated paper."""
    llm = ChatGoogleGenerativeAI(
        model="gemini-3.6-flash", 
        api_key=state["api_key"]
    )
    
    prompt_text = f"""
    You are an expert teacher. You just created a question paper for {state['subject']} (Class: {state['grade_class']}).
    Here is the question paper:
    
    {state['generated_paper']}
    
    Please generate a highly detailed answer key and marking scheme for the above question paper.
    Format the output clearly.
    """
    
    message = HumanMessage(content=[{"type": "text", "text": prompt_text}])
    response = llm.invoke([message])
    
    content = response.content
    if isinstance(content, list):
        text_parts = []
        for item in content:
            if isinstance(item, dict) and "text" in item:
                text_parts.append(item["text"])
            elif isinstance(item, str):
                text_parts.append(item)
        text_content = "\n".join(text_parts)
    else:
        text_content = str(content)
        
    return {"generated_answer_key": text_content}

# ==========================================
# 3. Build the Workflow Graph
# ==========================================
workflow = StateGraph(PaperState)
workflow.add_node("generator", generate_question_paper)
workflow.add_node("answer_key_generator", generate_answer_key)

workflow.set_entry_point("generator")
workflow.add_edge("generator", "answer_key_generator")
workflow.add_edge("answer_key_generator", END)
app_graph = workflow.compile()

# ==========================================
# Helper Function
# ==========================================
@st.cache_data(show_spinner=False)
def fetch_ncert_chapters(api_key, subject, grade_class):
    if not api_key or not subject or not grade_class:
        return []
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-3.6-flash")
        prompt = f"List all the exact NCERT chapters for Class {grade_class} {subject}. Return ONLY a raw JSON array of strings in the format ['1. Chapter Name', '2. Chapter Name']. Do not include any markdown formatting, backticks, or code blocks. Just the JSON array."
        response = model.generate_content(prompt)
        import json
        text = response.text.strip()
        if text.startswith('```json'):
            text = text[7:-3].strip()
        elif text.startswith('```'):
            text = text[3:-3].strip()
        chapters = json.loads(text)
        return chapters
    except Exception as e:
        return []

# ==========================================
# 4. Streamlit UI
# ==========================================
st.set_page_config(page_title="AI Question Paper Generator", layout="wide")
st.title("📄 AI Question Paper Generator")

# Sidebar for API Key
with st.sidebar:
    st.header("Authentication")
    user_api_key = st.text_input("Enter your Google Gemini API Key:", type="password")
    st.caption("Get your key from [Google AI Studio](https://aistudio.google.com/).")

tab1, tab2 = st.tabs(["Generate Paper", "Past Papers Archive"])
with tab1:
    # Institution Details
    st.subheader("Institution Details")
    col_s1, col_s2, col_s3 = st.columns(3)
    with col_s1:
        school_name = st.text_input("School Name", value="Neha Public School",placeholder="e.g. ABC Public School")
    with col_s2:
        session_year = st.selectbox("Session", ["24-25", "25-26", "26-27", "27-28"], index=2)
    with col_s3:
        exam_time = st.selectbox("Time Allowed", ["1 Hour", "1.5 Hours", "2 Hours", "2.5 Hours", "3 Hours"], index=3)
    
    st.markdown("---")
    # Main Content
    col1, col2, col3 = st.columns(3)
    
    with col1:
        subject = st.text_input("Subject (e.g., Physics, History)")
    with col2:
        grade_class = st.text_input("Class/Grade (e.g., 10th, University)")
    with col3:
        marks = st.number_input("Total Marks", min_value=1, max_value=200, value=50)
        
    is_ncert = st.checkbox("Make Paper From Ncert syllabus")
    
    ncert_chapters = ""
    if is_ncert:
        if not user_api_key:
            st.warning("Please enter your API Key in the sidebar to fetch chapters automatically.")
            ncert_chapters = st.text_input("Enter Chapter Numbers (e.g., 1, 3, 4)")
        elif not subject or not grade_class:
            st.warning("Please enter the Subject and Class above to fetch chapters automatically.")
            ncert_chapters = st.text_input("Enter Chapter Numbers (e.g., 1, 3, 4)")
        else:
            with st.spinner("Fetching NCERT syllabus..."):
                chapters_list = fetch_ncert_chapters(user_api_key, subject, grade_class)
            
            if chapters_list:
                selected_chapters = st.multiselect("Select Chapters for the paper", chapters_list)
                ncert_chapters = ", ".join(selected_chapters)
            else:
                ncert_chapters = st.text_input("Enter Chapter Numbers manually (unable to auto-fetch)")
    
    uploaded_files = st.file_uploader(
        "Upload Source Material (Images & PDFs)", 
        type=["png", "jpg", "jpeg", "pdf"], 
        accept_multiple_files=True
    )
    
    board_format = st.checkbox("Use Haryana Board Format (HBSE)")
    
    comments = st.text_area("Optional Comments/Specific Instructions", placeholder="e.g., Make the questions high difficulty, focus on application-based concepts...")
    
    st.markdown("---")
    submit_button = st.button("Generate Question Paper")
    
    # ==========================================
    # 5. Execution Logic
    # ==========================================
    if submit_button:
        if not user_api_key:
            st.error("Please enter your Gemini API Key in the sidebar.")
        elif not subject or not grade_class:
            st.error("Please fill in the Subject and Class fields.")
        else:
            with st.spinner("Analyzing context and generating paper..."):
                
                # Convert uploaded files for the API (PDFs need File API)
                file_data_list = []
                if uploaded_files:
                    genai.configure(api_key=user_api_key)
                    for file in uploaded_files:
                        if file.type == "application/pdf":
                            # Upload PDF to Google GenAI File API
                            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                                tmp.write(file.read())
                                tmp_path = tmp.name
                            
                            uploaded_gfile = genai.upload_file(tmp_path, mime_type="application/pdf")
                            file_data_list.append({
                                "mime_type": file.type,
                                "file_uri": uploaded_gfile.uri
                            })
                            os.remove(tmp_path)
                        else:
                            file_data_list.append({
                                "mime_type": file.type,
                                "data": base64.b64encode(file.read()).decode("utf-8")
                            })
                
                # Prepare inputs for LangGraph
                initial_state = {
                    "api_key": user_api_key,
                    "marks": marks,
                    "subject": subject,
                    "grade_class": grade_class,
                    "school_name": school_name,
                    "session": session_year,
                    "exam_time": exam_time,
                    "board_format": board_format,
                    "is_ncert": is_ncert,
                    "ncert_chapters": ncert_chapters,
                    "comments": comments,
                    "file_data": file_data_list
                }
                
                # Run the graph
                try:
                    result = app_graph.invoke(initial_state)
                    
                    st.success("Paper Generated Successfully!")
                    st.markdown("---")
                    st.markdown(result["generated_paper"])
                    
                    # Create a Word Document from the Markdown
                    def create_docx(text, state_data):
                        doc = Document()
                        
                        # Add Header Information
                        if state_data.get("school_name"):
                            p = doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run(state_data["school_name"])
                            run.bold = True
                            run.font.size = Pt(16)
                        
                        if state_data.get("session"):
                            p = doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p.add_run(f"Session {state_data['session']}")
                            
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        header_text = f"Class: {state_data['grade_class']} | Subject: {state_data['subject']} | Time: {state_data['exam_time']} | Max Marks: {state_data['marks']}"
                        run = p.add_run(header_text)
                        run.bold = True
                        
                        doc.add_paragraph("-" * 50)
                        
                        import markdown
                        from htmldocx import HtmlToDocx
                        
                        html_content = markdown.markdown(text, extensions=['tables', 'fenced_code'])
                        parser = HtmlToDocx()
                        parser.add_html_to_document(html_content, doc)
                        buffer = io.BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        return buffer
    
                    paper_docx_buffer = create_docx(result["generated_paper"], initial_state)
                    answer_key_docx_buffer = create_docx(result["generated_answer_key"], initial_state)
                    
                    # Save to archives
                    os.makedirs("archives", exist_ok=True)
                    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                    paper_filename = f"archives/paper_{timestamp}.docx"
                    answer_filename = f"archives/answer_{timestamp}.docx"
                    
                    with open(paper_filename, "wb") as f:
                        f.write(paper_docx_buffer.getvalue())
                    with open(answer_filename, "wb") as f:
                        f.write(answer_key_docx_buffer.getvalue())
                        
                    # Update history.json
                    history_file = "archives/history.json"
                    if os.path.exists(history_file):
                        with open(history_file, "r") as f:
                            history = json.load(f)
                    else:
                        history = []
                        
                    history.append({
                        "id": timestamp,
                        "date": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
                        "class": grade_class,
                        "subject": subject,
                        "marks": marks,
                        "paper_file": paper_filename,
                        "answer_file": answer_filename
                    })
                    
                    with open(history_file, "w") as f:
                        json.dump(history, f, indent=4)
                    
                    # Add download buttons
                    col_btn1, col_btn2 = st.columns(2)
                    with col_btn1:
                        st.download_button(
                            label="Download Paper as Word (.docx)",
                            data=paper_docx_buffer,
                            file_name=f"{subject}_paper.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    with col_btn2:
                        st.download_button(
                            label="Download Answer Key (.docx)",
                            data=answer_key_docx_buffer,
                            file_name=f"ANSWER KEY {grade_class} {subject}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                except Exception as e:
                    st.error(f"An error occurred: {e}")

with tab2:
    st.header("Past Papers Archive")
    history_file = "archives/history.json"
    
    if os.path.exists(history_file):
        with open(history_file, "r") as f:
            history = json.load(f)
            
        if not history:
            st.info("No past papers found.")
        else:
            # Group by class then subject
            grouped = {}
            for entry in history:
                c = entry.get("class", "Unknown")
                s = entry.get("subject", "Unknown")
                if c not in grouped:
                    grouped[c] = {}
                if s not in grouped[c]:
                    grouped[c][s] = []
                grouped[c][s].append(entry)
                
            for c in sorted(grouped.keys()):
                st.subheader(f"Class: {c}")
                for s in sorted(grouped[c].keys()):
                    with st.expander(f"Subject: {s} ({len(grouped[c][s])} papers)"):
                        for entry in reversed(grouped[c][s]):
                            st.markdown(f"**Date:** {entry.get('date')} | **Marks:** {entry.get('marks')}")
                            
                            col_dl1, col_dl2 = st.columns(2)
                            try:
                                with open(entry['paper_file'], "rb") as pf:
                                    col_dl1.download_button("Download Paper", data=pf, file_name=f"{s}_paper_{entry['id']}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"p_{entry['id']}")
                                with open(entry['answer_file'], "rb") as af:
                                    col_dl2.download_button("Download Answer Key", data=af, file_name=f"ANSWER_KEY_{c}_{s}_{entry['id']}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"a_{entry['id']}")
                            except FileNotFoundError:
                                st.error("Files for this entry are missing from storage.")
                            st.divider()
    else:
        st.info("No past papers found yet. Generate a paper to start your archive!")
